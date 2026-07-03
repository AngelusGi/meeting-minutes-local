"""
Minute Agent - trascrizione locale di meeting e generazione minute.

Pipeline: video OBS -> ffmpeg (estrazione audio) -> faster-whisper (trascrizione)
          -> LLM locale via Ollama (minuta) -> export .md + .docx

Default: tutto in locale. L'architettura a engine permette di aggiungere
provider cloud (es. Azure Speech / MAI-Transcribe) implementando una nuova
classe TranscriptionEngine.
"""

import json
import re
import shutil
import subprocess
import tempfile
import threading
import time
import traceback
import uuid
from pathlib import Path

import urllib.request
import urllib.error

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

BASE_DIR = Path(__file__).resolve().parent
JOBS_DIR = BASE_DIR / "jobs"
STATIC_DIR = BASE_DIR / "static"
JOBS_DIR.mkdir(exist_ok=True)

OLLAMA_URL = "http://localhost:11434"
DEFAULT_OLLAMA_MODEL = "llama3.1:8b"
CONFIG_PATH = BASE_DIR / "config.json"

VALID_WHISPER_MODELS = ["tiny", "base", "small", "medium", "large-v3", "large-v3-turbo"]
VALID_LANGUAGES = ["auto", "it", "en"]
VALID_AZURE_MODELS = ["mai-transcribe-1.5", "mai-transcribe-1"]


def load_config() -> dict:
    try:
        return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}


def save_config(cfg: dict):
    CONFIG_PATH.write_text(json.dumps(cfg, indent=2), encoding="utf-8")

app = FastAPI(title="Minute Agent")

# ----------------------------------------------------------------------------
# Job store (in memory + su disco per i file)
# ----------------------------------------------------------------------------

JOBS: dict = {}
JOBS_LOCK = threading.Lock()


def update_job(job_id: str, **kwargs):
    with JOBS_LOCK:
        JOBS[job_id].update(kwargs)


def append_log(job_id: str, msg: str):
    line = f"[{time.strftime('%H:%M:%S')}] {msg}"
    print(line, flush=True)  # visibile anche nella console del server
    with JOBS_LOCK:
        JOBS[job_id]["log"].append(line)


# ----------------------------------------------------------------------------
# Step 1: estrazione audio con ffmpeg
# ----------------------------------------------------------------------------

def ffmpeg_path() -> str:
    """ffmpeg di sistema se presente, altrimenti binario statico bundled
    via imageio-ffmpeg (nessuna installazione di sistema richiesta)."""
    system = shutil.which("ffmpeg")
    if system:
        return system
    try:
        import imageio_ffmpeg
        return imageio_ffmpeg.get_ffmpeg_exe()
    except Exception:
        raise RuntimeError(
            "ffmpeg non disponibile: ne' di sistema ne' bundled. "
            "Reinstalla le dipendenze (pip install -r requirements.txt)."
        )


def extract_audio(video_path: Path, out_dir: Path, fmt: str = "wav") -> Path:
    """wav 16k mono per la trascrizione locale; flac (compresso lossless)
    per il cloud, che ha un limite di 300 MB per file."""
    if fmt == "flac":
        audio_path = out_dir / "audio.flac"
        codec = ["-c:a", "flac"]
    else:
        audio_path = out_dir / "audio.wav"
        codec = ["-c:a", "pcm_s16le"]
    cmd = [
        ffmpeg_path(), "-y", "-i", str(video_path),
        "-vn", "-ac", "1", "-ar", "16000", *codec,
        str(audio_path),
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0:
        raise RuntimeError(f"ffmpeg errore: {proc.stderr[-800:]}")
    return audio_path


# ----------------------------------------------------------------------------
# Step 2: trascrizione (architettura a engine, locale di default)
# ----------------------------------------------------------------------------

class TranscriptionEngine:
    """Interfaccia base. Per aggiungere un provider cloud (es. Azure Speech /
    MAI-Transcribe) creare una sottoclasse e registrarla in ENGINES."""

    def transcribe(self, audio_path: Path, language: str, job_id: str) -> dict:
        raise NotImplementedError


class FasterWhisperEngine(TranscriptionEngine):
    def __init__(self, model_size: str):
        self.model_size = model_size

    def transcribe(self, audio_path: Path, language: str, job_id: str) -> dict:
        from faster_whisper import WhisperModel

        device, compute_type = self._pick_device()
        append_log(job_id, f"Carico faster-whisper '{self.model_size}' su {device} ({compute_type})")
        append_log(job_id, "Se il modello non e' ancora sul PC viene scaricato ora: "
                           "progresso visibile nella console del server")
        try:
            model = WhisperModel(self.model_size, device=device, compute_type=compute_type)
        except Exception:
            if device == "cuda":
                append_log(job_id, "GPU non utilizzabile, passo a CPU (int8)")
                model = WhisperModel(self.model_size, device="cpu", compute_type="int8")
            else:
                raise

        lang = None if language == "auto" else language
        segments_iter, info = model.transcribe(
            str(audio_path), language=lang, vad_filter=True, beam_size=5
        )

        segments = []
        text_parts = []
        total = info.duration or 0
        append_log(job_id, f"Audio: {_human_duration(total)}. Inizio trascrizione...")
        last_beat = time.time()
        pct = 0
        for seg in segments_iter:
            segments.append({"start": seg.start, "end": seg.end, "text": seg.text.strip()})
            text_parts.append(seg.text.strip())
            if total:
                pct = min(99, int(seg.end / total * 100))
                update_job(job_id, step_progress=pct)
            if time.time() - last_beat > 20:
                append_log(job_id, f"Trascrizione in corso: {_human_duration(seg.end)} "
                                   f"di {_human_duration(total)} ({pct}%)")
                last_beat = time.time()

        return {
            "text": "\n".join(text_parts),
            "segments": segments,
            "language": info.language,
            "duration": info.duration,
        }

    @staticmethod
    def _pick_device():
        try:
            import ctranslate2
            if ctranslate2.get_cuda_device_count() > 0:
                return "cuda", "float16"
        except Exception:
            pass
        return "cpu", "int8"


class AzureFoundryEngine(TranscriptionEngine):
    """MAI-Transcribe via Azure AI Foundry (LLM Speech API, preview).
    POST multipart a /speechtotext/transcriptions:transcribe."""

    API_VERSION = "2025-10-15"
    MAX_BYTES = 300 * 1024 * 1024  # limite documentato del servizio

    def __init__(self, endpoint: str, key: str, model: str):
        endpoint = (endpoint or "").strip().rstrip("/")
        if endpoint and not endpoint.startswith("http"):
            # accetta anche il solo nome della risorsa Foundry/Speech
            endpoint = f"https://{endpoint}.cognitiveservices.azure.com"
        if not endpoint or not key:
            raise RuntimeError(
                "Cloud Azure AI Foundry non configurato: servono endpoint "
                "(es. nome-risorsa oppure https://nome-risorsa.cognitiveservices.azure.com) "
                "e chiave della risorsa Speech/Foundry."
            )
        self.endpoint, self.key = endpoint, key
        # testo libero: qualsiasi modello supportato dalla LLM Speech API
        self.model = (model or "").strip() or VALID_AZURE_MODELS[0]

    def transcribe(self, audio_path: Path, language: str, job_id: str) -> dict:
        import requests

        size = audio_path.stat().st_size
        if size > self.MAX_BYTES:
            raise RuntimeError(
                f"Audio {size // 1048576} MB: oltre il limite di 300 MB "
                "dell'API. Usa la trascrizione locale per registrazioni cosi' lunghe."
            )
        definition = {"enhancedMode": {"enabled": True, "model": self.model}}
        if language != "auto":
            definition["locales"] = [language]

        url = f"{self.endpoint}/speechtotext/transcriptions:transcribe?api-version={self.API_VERSION}"
        append_log(job_id, f"Invio audio ({size // 1048576} MB) a Azure AI Foundry "
                           f"({self.model}) - attendi la risposta del servizio")
        with open(audio_path, "rb") as f:
            resp = requests.post(
                url,
                headers={"Ocp-Apim-Subscription-Key": self.key},
                files={
                    "audio": (audio_path.name, f, "audio/flac"),
                    "definition": (None, json.dumps(definition), "application/json"),
                },
                timeout=3600,
            )
        if resp.status_code != 200:
            raise RuntimeError(
                f"Azure AI Foundry ha risposto {resp.status_code}: {resp.text[:500]}"
            )
        data = resp.json()
        phrases = data.get("phrases", [])
        segments = [
            {
                "start": p.get("offsetMilliseconds", 0) / 1000,
                "end": (p.get("offsetMilliseconds", 0) + p.get("durationMilliseconds", 0)) / 1000,
                "text": p.get("text", "").strip(),
            }
            for p in phrases
        ]
        text = "\n".join(cp.get("text", "") for cp in data.get("combinedPhrases", []))
        detected = phrases[0].get("locale", "auto") if phrases else "auto"
        append_log(job_id, "Risposta ricevuta da Azure AI Foundry")
        return {
            "text": text,
            "segments": segments,
            "language": detected,
            "duration": data.get("durationMilliseconds", 0) / 1000,
        }


def get_engine(computation: str, model_size: str, azure_model: str = "") -> TranscriptionEngine:
    if computation == "local":
        return FasterWhisperEngine(model_size)
    cfg = load_config().get("azure", {})
    return AzureFoundryEngine(cfg.get("endpoint", ""), cfg.get("key", ""),
                              azure_model or cfg.get("model", ""))


# ----------------------------------------------------------------------------
# Step 3: generazione minuta (LLM locale via Ollama, con fallback template)
# ----------------------------------------------------------------------------

MINUTE_PROMPT = """Sei un consulente IT senior. Dalla trascrizione di un meeting di consulenza IT, scrivi una minuta professionale e sintetica NELLA STESSA LINGUA della trascrizione.

Struttura obbligatoria (markdown):

# Minuta meeting - {title}

**Data:** {date}
**Durata:** {duration}

## Partecipanti
(OBBLIGATORIO: individua sempre gli interlocutori dalla trascrizione - nomi propri pronunciati, presentazioni, riferimenti diretti, ruoli e azienda se deducibili. Elencali con eventuale ruolo. Se qualcuno non e' identificabile, indicalo come "Interlocutore non identificato" e aggiungi "da completare")

## Sintesi esecutiva
(3-5 frasi)

## Punti discussi
(elenco puntato, raggruppato per tema)

## Decisioni prese
(elenco; se nessuna, scrivi "Nessuna decisione formale")

## Azioni (action item)
| Azione | Owner | Scadenza |
|---|---|---|

## Rischi e punti aperti
(elenco)

## Prossimi passi
(elenco)

Regole: non inventare fatti non presenti nella trascrizione; usa "da completare" dove mancano informazioni; mantieni i termini tecnici originali (Azure, Terraform, policy, ecc.); attribuisci decisioni e azioni alle persone quando l'interlocutore e' identificabile dal contesto.
{custom}
TRASCRIZIONE:
{transcript}
"""


def generate_minutes(transcript: str, title: str, date: str, duration: str,
                     ollama_model: str, job_id: str, custom_prompt: str = "") -> tuple:
    """Ritorna (markdown, warning|None)."""
    custom = ""
    if custom_prompt.strip():
        custom = ("\nISTRUZIONI AGGIUNTIVE DELL'UTENTE (hanno priorita' sulle regole "
                  f"precedenti, tranne il divieto di inventare fatti):\n{custom_prompt.strip()}\n")
    prompt = MINUTE_PROMPT.format(
        title=title, date=date, duration=duration, custom=custom,
        transcript=transcript[:60000]
    )
    try:
        append_log(job_id, f"Genero la minuta con Ollama ({ollama_model}) - "
                           "su CPU puo' richiedere alcuni minuti, attendi")
        payload = json.dumps({
            "model": ollama_model,
            "messages": [{"role": "user", "content": prompt}],
            "stream": False,
            "options": {"temperature": 0.2},
        }).encode()
        req = urllib.request.Request(
            f"{OLLAMA_URL}/api/chat", data=payload,
            headers={"Content-Type": "application/json"},
        )
        with urllib.request.urlopen(req, timeout=1800) as resp:
            data = json.loads(resp.read())
        md = data["message"]["content"].strip()
        if not md.startswith("#"):
            md = f"# Minuta meeting - {title}\n\n" + md
        return md, None
    except Exception as e:
        warning = (
            f"Ollama non raggiungibile o errore ({e}). Ho prodotto un template "
            "di minuta con la trascrizione allegata: compila le sezioni a mano, "
            "oppure avvia Ollama (https://ollama.com) e rilancia."
        )
        append_log(job_id, warning)
        md = MINUTE_PROMPT.format(
            title=title, date=date, duration=duration, custom="", transcript=""
        ).split("TRASCRIZIONE:")[0].split("Regole:")[0]
        md = md.replace("Sei un consulente IT senior.", "").strip()
        md = md.split("Struttura obbligatoria (markdown):", 1)[-1].strip()
        md += "\n\n---\n\n## Trascrizione integrale\n\n" + transcript
        return md, warning


# ----------------------------------------------------------------------------
# Step 4: export docx (conversione markdown minimale)
# ----------------------------------------------------------------------------

def markdown_to_docx(md: str, out_path: Path):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue
        if line.startswith("### "):
            doc.add_heading(_strip_md(line[4:]), level=3)
        elif line.startswith("## "):
            doc.add_heading(_strip_md(line[3:]), level=2)
        elif line.startswith("# "):
            doc.add_heading(_strip_md(line[2:]), level=1)
        elif re.match(r"^\s*[-*] ", line):
            _add_runs(doc.add_paragraph(style="List Bullet"), re.sub(r"^\s*[-*] ", "", line))
        elif re.match(r"^\s*\d+\. ", line):
            _add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\s*\d+\. ", "", line))
        elif line.strip() == "---":
            doc.add_paragraph()
        else:
            _add_runs(doc.add_paragraph(), line)
        i += 1
    doc.save(str(out_path))


def _strip_md(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text).strip()


def _add_runs(paragraph, text: str):
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*", text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_table(doc, table_lines):
    rows = []
    for tl in table_lines:
        if re.match(r"^\|[\s:-]+\|", tl.replace("-", "-")) and set(tl) <= set("|-: "):
            continue
        cells = [c.strip() for c in tl.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c in range(ncols):
            table.cell(r, c).text = _strip_md(row[c]) if c < len(row) else ""


# ----------------------------------------------------------------------------
# Pipeline
# ----------------------------------------------------------------------------

def run_pipeline(job_id: str, video_path: Path, opts: dict):
    job_dir = JOBS_DIR / job_id
    try:
        update_job(job_id, status="running", step="audio", step_progress=0)
        append_log(job_id, "Estraggo l'audio dal video (ffmpeg)")
        fmt = "flac" if opts["computation"] == "cloud" else "wav"
        audio_path = extract_audio(video_path, job_dir, fmt)
        update_job(job_id, step="transcribe", step_progress=0)

        engine = get_engine(opts["computation"], opts["model_size"], opts.get("azure_model", ""))
        result = engine.transcribe(audio_path, opts["language"], job_id)

        transcript_path = job_dir / "trascrizione.txt"
        transcript_path.write_text(result["text"], encoding="utf-8")
        srt_lines = []
        for n, seg in enumerate(result["segments"], 1):
            srt_lines.append(f"{n}\n{_ts(seg['start'])} --> {_ts(seg['end'])}\n{seg['text']}\n")
        (job_dir / "trascrizione.srt").write_text("\n".join(srt_lines), encoding="utf-8")
        append_log(job_id, f"Trascrizione completata (lingua: {result['language']})")

        update_job(job_id, step="minutes", step_progress=0)
        duration_str = _human_duration(result.get("duration") or 0)
        md, warning = generate_minutes(
            result["text"], opts["title"], opts["date"], duration_str,
            opts["ollama_model"], job_id, opts.get("custom_prompt", ""),
        )
        (job_dir / "minuta.md").write_text(md, encoding="utf-8")

        update_job(job_id, step="export", step_progress=50)
        markdown_to_docx(md, job_dir / "minuta.docx")
        append_log(job_id, "Export .md e .docx completato")

        update_job(
            job_id, status="done", step="done", step_progress=100,
            warning=warning,
            outputs=["minuta.md", "minuta.docx", "trascrizione.txt", "trascrizione.srt"],
            detected_language=result["language"],
        )
    except Exception as e:
        traceback.print_exc()
        update_job(job_id, status="error", error=str(e))
        append_log(job_id, f"ERRORE: {e}")
    finally:
        try:
            video_path.unlink(missing_ok=True)
            (job_dir / "audio.wav").unlink(missing_ok=True)
            (job_dir / "audio.flac").unlink(missing_ok=True)
        except Exception:
            pass


def _ts(seconds: float) -> str:
    ms = int((seconds % 1) * 1000)
    s = int(seconds)
    return f"{s // 3600:02d}:{s % 3600 // 60:02d}:{s % 60:02d},{ms:03d}"


def _human_duration(seconds: float) -> str:
    m = int(seconds // 60)
    return f"{m // 60}h {m % 60}min" if m >= 60 else f"{m} min"


# ----------------------------------------------------------------------------
# API
# ----------------------------------------------------------------------------

@app.post("/api/jobs")
async def create_job(
    file: UploadFile = File(...),
    model_size: str = Form("small"),
    language: str = Form("auto"),
    computation: str = Form("local"),
    title: str = Form("Meeting"),
    date: str = Form(""),
    ollama_model: str = Form(DEFAULT_OLLAMA_MODEL),
    azure_endpoint: str = Form(""),
    azure_key: str = Form(""),
    azure_model: str = Form(""),
    custom_prompt: str = Form(""),
):
    if model_size not in VALID_WHISPER_MODELS:
        raise HTTPException(400, f"Modello non valido: {model_size}")
    if language not in VALID_LANGUAGES:
        raise HTTPException(400, f"Lingua non valida: {language}")

    # persisti il prompt personalizzato (config.json locale)
    cfgp = load_config()
    if cfgp.get("custom_prompt", "") != custom_prompt:
        cfgp["custom_prompt"] = custom_prompt
        save_config(cfgp)

    # persisti la configurazione Azure se fornita (config.json locale)
    if computation == "cloud":
        cfg = load_config()
        az = cfg.setdefault("azure", {})
        if azure_endpoint.strip():
            az["endpoint"] = azure_endpoint.strip()
        if azure_key.strip():
            az["key"] = azure_key.strip()
        if azure_model.strip():
            az["model"] = azure_model.strip()
        save_config(cfg)
        if not az.get("endpoint") or not az.get("key"):
            raise HTTPException(400, "Cloud selezionato ma endpoint o chiave Azure mancanti.")

    job_id = uuid.uuid4().hex[:12]
    job_dir = JOBS_DIR / job_id
    job_dir.mkdir(parents=True)

    suffix = Path(file.filename or "video.mkv").suffix or ".mkv"
    video_path = job_dir / f"input{suffix}"
    with open(video_path, "wb") as f:
        while chunk := await file.read(1024 * 1024):
            f.write(chunk)

    opts = {
        "model_size": model_size,
        "language": language,
        "computation": computation,
        "title": title.strip() or "Meeting",
        "date": date.strip() or time.strftime("%d/%m/%Y"),
        "ollama_model": ollama_model.strip() or DEFAULT_OLLAMA_MODEL,
        "azure_model": azure_model,
        "custom_prompt": custom_prompt,
    }
    with JOBS_LOCK:
        JOBS[job_id] = {
            "id": job_id, "status": "queued", "step": "queued",
            "step_progress": 0, "log": [], "opts": opts,
            "filename": file.filename, "outputs": [], "warning": None, "error": None,
        }
    threading.Thread(target=run_pipeline, args=(job_id, video_path, opts), daemon=True).start()
    return {"job_id": job_id}


@app.get("/api/jobs/{job_id}")
def job_status(job_id: str):
    with JOBS_LOCK:
        job = JOBS.get(job_id)
        if not job:
            raise HTTPException(404, "Job non trovato")
        return JSONResponse(dict(job))


@app.get("/api/jobs/{job_id}/download/{filename}")
def download(job_id: str, filename: str):
    safe = Path(filename).name
    path = JOBS_DIR / job_id / safe
    if not path.exists():
        raise HTTPException(404, "File non trovato")
    return FileResponse(str(path), filename=safe)


@app.get("/api/config")
def get_config():
    az = load_config().get("azure", {})
    return {
        "azure_endpoint": az.get("endpoint", ""),
        "azure_model": az.get("model", VALID_AZURE_MODELS[0]),
        "azure_has_key": bool(az.get("key")),
        "custom_prompt": load_config().get("custom_prompt", ""),
    }


@app.get("/api/health")
def health():
    try:
        ffmpeg_ok = bool(ffmpeg_path())
    except RuntimeError:
        ffmpeg_ok = False
    ollama_ok = False
    try:
        with urllib.request.urlopen(f"{OLLAMA_URL}/api/tags", timeout=2) as resp:
            ollama_ok = resp.status == 200
    except Exception:
        pass
    gpu = False
    try:
        import ctranslate2
        gpu = ctranslate2.get_cuda_device_count() > 0
    except Exception:
        pass
    return {"ffmpeg": ffmpeg_ok, "ollama": ollama_ok, "gpu": gpu}


@app.get("/", response_class=HTMLResponse)
def index():
    return (STATIC_DIR / "index.html").read_text(encoding="utf-8")


app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8756)
